from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
import pandas as pd
from openpyxl.styles import Alignment, Border, Side
import datetime

class WeeklyReportGenerator:
    def __init__(self, year=datetime.datetime.now().year, 
                 week=datetime.date.today().isocalendar()[1], 
                 month=datetime.datetime.now().month, 
                 day_start=datetime.datetime.now().day-4, 
                 day_end=datetime.datetime.now().day, 
                 name='杨佳'):
        self.year = year
        self.week = week
        self.month = month
        self.day_start = day_start
        self.day_end = day_end
        self.name = name
        self.doc = Document()
        self.setup_document()

    def setup_document(self):
        self.doc_name = f'系统与测试部AI开发部工作周报-{self.year}W{self.week}-{self.name}'
        self.greeting_str = f'戴总，宋哥，各位同事好！\n        这是{self.name}本周周报，请查收。'
        self.report_str = f'报告时间：{self.year}年{self.month}月{self.day_start}日—{self.year}年{self.month}月{self.day_end}日\n报告人： {self.name}'

    def add_greeting(self):
        greeting = self.doc.add_paragraph()
        greeting_runs = greeting.add_run(self.greeting_str)
        greeting_runs.font.size = Pt(12)

    def add_heading(self):
        heading = self.doc.add_paragraph()
        heading_runs = heading.add_run(self.doc_name)
        heading_runs.font.color.rgb = RGBColor(0,0,0)
        heading_runs.font.bold = True
        heading_runs.font.size = Pt(18)
        heading.paragraph_format.space_before = Pt(30)
        self.doc.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_report_info(self):
        report = self.doc.add_paragraph()
        report.add_run(self.report_str)
        report.runs[0].font.size = Pt(12)

    def add_main_work(self, project_names, jobs_list):
        main_work = self.doc.add_paragraph()
        main_work.add_run('一、本周主要工作内容')
        main_work.runs[0].font.size = Pt(12)
        main_work.runs[0].font.bold = True
        for i, project_name in enumerate(project_names):
            project_name_runs = self.doc.add_paragraph().add_run(str(i+1) + '.' + project_name)
            project_name_runs.font.bold = True
            project_name_runs.font.color.rgb = RGBColor(0,0,255)
            jobs = self.doc.add_paragraph()
            for job in jobs_list[i]:
                jobs.add_run(job + '\n')

    def add_questions(self, questions_list):
        questions = self.doc.add_paragraph()
        questions.add_run('二、存在问题')
        questions.runs[0].font.size = Pt(12)
        questions.runs[0].font.bold = True
        qss = self.doc.add_paragraph()
        for qs in questions_list:
            qss.add_run(qs + '\n')

    def add_next_week_plan(self, plans_list):
        new_work = self.doc.add_paragraph()
        new_work.add_run('三、下周重点工作计划')
        new_work.runs[0].font.size = Pt(12)
        new_work.runs[0].font.bold = True
        works = self.doc.add_paragraph()
        for work in plans_list:
            works.add_run(work + '\n')

    def set_font_style(self):
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def save_document(self, filename="test.docx"):
        self.doc.save(filename)
    
    def generate_excel(self, project_names, jobs_list, questions_list, plans_list):
        # 创建一个DataFrame来存储所有内容
        data = {
            '项目名称': project_names,
            '任务详情': ['\n'.join(jobs) for jobs in jobs_list],
            '存在问题': ['\n'.join(questions_list)] * len(project_names),
            '下周计划': ['\n'.join(plans_list)] * len(project_names)
        }
        df = pd.DataFrame(data)
        
        # 创建一个新的Excel工作簿
        with pd.ExcelWriter(self.doc_name + '.xlsx', engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='周报内容')
            ws = writer.sheets['周报内容']
            
            # 设置任务详情列的宽度为100像素
            ws.column_dimensions['A'].width = 20
            ws.column_dimensions['B'].width = 100
            ws.column_dimensions['C'].width = 50
            ws.column_dimensions['D'].width = 50
            
            # 设置自动换行
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=4):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True)
            
            # 合并存在问题列的单元格
            ws.merge_cells(start_row=2, start_column=3, end_row=ws.max_row, end_column=3)
            ws.cell(row=2, column=3).alignment = Alignment(wrap_text=True, vertical='top')
            
            # 合并下周计划列的单元格
            ws.merge_cells(start_row=2, start_column=4, end_row=ws.max_row, end_column=4)
            ws.cell(row=2, column=4).alignment = Alignment(wrap_text=True, vertical='top')

            # 添加表格框线
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=4):
                for cell in row:
                    cell.border = thin_border
        
        print(f"Excel文件已保存为 {self.doc_name}.xlsx")

    def generate_report(self, project_names, jobs_list, questions_list, plans_list):
        self.add_greeting()
        self.add_heading()
        self.add_report_info()
        self.add_main_work(project_names, jobs_list)
        self.add_questions(questions_list)
        self.add_next_week_plan(plans_list)
        self.set_font_style()
        self.save_document(self.doc_name + '.docx')
        self.generate_excel(project_names, jobs_list, questions_list, plans_list)
        self.doc = Document()

    def generate_html_report(self, project_names, jobs_list, questions_list, plans_list):
        greeting = self.greeting_str.replace('\n', '<br>\t')
        report = self.report_str.replace('\n', '<br>')
        html_content = f"""
        <html>
        <body>
        <p>{greeting}</p>
        <h1 style="text-align: center;">{self.doc_name}</h1>
        <p>{report}</p>
        <h2>一、本周主要工作内容</h2>
        """
        for i, project_name in enumerate(project_names):
            html_content += f"<h3 style='color: blue;'>{project_name}</h3><ul>"
            for job in jobs_list[i]:
                html_content += f"<li>{job}</li>"
            html_content += "</ul>"
        
        html_content += "<h2>二、存在问题</h2><ul>"
        for qs in questions_list:
            html_content += f"<li>{qs}</li>"
        html_content += "</ul>"

        html_content += "<h2>三、下周重点工作计划</h2><ul>"
        for work in plans_list:
            html_content += f"<li>{work}</li>"
        html_content += "</ul></body></html>"

        return html_content

    @staticmethod
    def read_report(filename):
        doc = Document(filename)
        content = {
            "greeting": "",
            "heading": "",
            "report_info": "",
            "main_work": {},
            "questions": [],
            "next_week_plan": []
        }
        
        paragraphs = list(doc.paragraphs)
        if len(paragraphs) > 0:
            content["greeting"] = paragraphs[0].text
        if len(paragraphs) > 1:
            content["heading"] = paragraphs[1].text
        if len(paragraphs) > 2:
            content["report_info"] = paragraphs[2].text
        
        main_work_start = next((i for i, p in enumerate(paragraphs) if p.text.startswith('一、')), None)
        questions_start = next((i for i, p in enumerate(paragraphs) if p.text.startswith('二、')), None)
        if main_work_start and questions_start:
            main_work_section = paragraphs[main_work_start+1:questions_start]
            current_project = None
            for paragraph in main_work_section:
                if any(run.bold for run in paragraph.runs):  # Check if any run in the paragraph is bold
                    current_project = paragraph.text
                    content["main_work"][current_project] = []
                elif current_project:  # Ensure there is a current project set before adding tasks
                    content["main_work"][current_project].append(paragraph.text)
                else:
                    # Optionally handle or log the error when there is no project set
                    pass

        if questions_start:
            next_week_plan_start = next((i for i, p in enumerate(paragraphs) if p.text.startswith('三、')), None)
            content["questions"] = [p.text for p in paragraphs[questions_start+1:next_week_plan_start]]

        if next_week_plan_start:
            content["next_week_plan"] = [p.text for p in paragraphs[next_week_plan_start+1:]]

        return content


# 使用示例
if __name__ == "__main__":
    report_generator = WeeklyReportGenerator(year=2024, week=17, month=4, day_start=22, day_end=26, name='杨佳')
    project_names = ['项目一名称', '项目二名称']
    jobs_list = [
        ['任务1详情', '任务2详情'],
        ['任务1详情', '任务2详情', '任务3详情']
    ]
    questions_list = ['问题1详情', '问题2详情']
    plans_list = ['下周计划1详情', '下周计划2详情']

    report_generator.generate_report(
        project_names=project_names,
        jobs_list=jobs_list,
        questions_list=questions_list,
        plans_list=plans_list
    )
    html_body = report_generator.generate_html_report(
        project_names=project_names,
        jobs_list=jobs_list,
        questions_list=questions_list,
        plans_list=plans_list
    )
    print(html_body)
    # filename = "系统与测试部AI开发部工作周报-2024W17-杨佳.docx"
    # print(WeeklyReportGenerator.read_report(filename))